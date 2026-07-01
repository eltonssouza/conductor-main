"""`cdt doc` — render a Markdown file to a .docx and/or .pdf deliverable.

The Conductor flow writes specs and client-questions as Markdown (the source of
truth, ingested into the journal). This turns one of those into an editable
client deliverable — deterministically, in the CLI, so it works the same on any
harness/LLM (including immature ones that cannot generate documents themselves).

Markdown is the source; the rendered doc is a derived artifact. By default it
lands in `<project>/.cdt/deliverables/` (outside the memory tree → never ingested).

Requires the optional extra: `pip install conductor[docs]`
(`python-docx` for .docx, `reportlab` for .pdf).
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Dict, List, Optional

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}.*$")


def _split_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


# --- markdown → block model (parsed once, rendered per format) ----------------

def parse_blocks(md: str) -> List[Dict]:
    """Parse a (small) Markdown subset into a flat list of block dicts:
    heading / para / bullet / number / code / table. Inline **bold** is kept as
    literal markup in the text and handled by each renderer."""
    blocks: List[Dict] = []
    lines = md.splitlines()
    i = 0
    in_code = False
    code_buf: List[str] = []
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                blocks.append({"type": "code", "text": "\n".join(code_buf)})
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.lstrip().startswith("|") and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1]):
            header = _split_row(line)
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                rows.append(_split_row(lines[j]))
                j += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            i = j
            continue

        m = _HEADING_RE.match(line)
        if m:
            blocks.append({"type": "heading", "level": min(len(m.group(1)), 4), "text": m.group(2)})
            i += 1
            continue
        m = _BULLET_RE.match(line)
        if m:
            blocks.append({"type": "bullet", "text": m.group(1)})
            i += 1
            continue
        m = _NUMBERED_RE.match(line)
        if m:
            blocks.append({"type": "number", "text": m.group(1)})
            i += 1
            continue
        if line.strip():
            blocks.append({"type": "para", "text": line.strip()})
        i += 1

    if code_buf:
        blocks.append({"type": "code", "text": "\n".join(code_buf)})
    return blocks


# --- .docx renderer (python-docx) --------------------------------------------

def _docx_runs(paragraph, text: str) -> None:
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_docx(blocks: List[Dict], out: Path, title: Optional[str]) -> None:
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for b in blocks:
        t = b["type"]
        if t == "heading":
            doc.add_heading(b["text"], level=b["level"])
        elif t == "bullet":
            _docx_runs(doc.add_paragraph(style="List Bullet"), b["text"])
        elif t == "number":
            _docx_runs(doc.add_paragraph(style="List Number"), b["text"])
        elif t == "code":
            doc.add_paragraph().add_run(b["text"]).font.name = "Consolas"
        elif t == "table":
            header, rows = b["header"], b["rows"]
            table = doc.add_table(rows=1, cols=len(header))
            try:
                table.style = "Light Grid Accent 1"
            except KeyError:
                pass
            for c, txt in enumerate(header):
                _docx_runs(table.rows[0].cells[c].paragraphs[0], txt)
            for r in rows:
                cells = table.add_row().cells
                for c in range(len(header)):
                    _docx_runs(cells[c].paragraphs[0], r[c] if c < len(r) else "")
        else:
            _docx_runs(doc.add_paragraph(), b["text"])
    doc.save(str(out))


# --- .pdf renderer (reportlab) -----------------------------------------------

def _pdf_inline(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return _BOLD_RE.sub(r"<b>\1</b>", text)


def render_pdf(blocks: List[Dict], out: Path, title: Optional[str]) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    flow = []
    if title:
        flow.append(Paragraph(_pdf_inline(title), styles["Title"]))
    for b in blocks:
        t = b["type"]
        if t == "heading":
            flow.append(Paragraph(_pdf_inline(b["text"]), styles[f"Heading{min(b['level'], 4)}"]))
        elif t in ("bullet", "number"):
            bt = "bullet" if t == "bullet" else "1"
            flow.append(ListFlowable(
                [ListItem(Paragraph(_pdf_inline(b["text"]), styles["BodyText"]))],
                bulletType=bt))
        elif t == "code":
            flow.append(Paragraph(_pdf_inline(b["text"]).replace("\n", "<br/>"), styles["Code"]))
        elif t == "table":
            data = [[Paragraph(_pdf_inline(x), styles["BodyText"]) for x in b["header"]]]
            for r in b["rows"]:
                row = [r[c] if c < len(r) else "" for c in range(len(b["header"]))]
                data.append([Paragraph(_pdf_inline(x), styles["BodyText"]) for x in row])
            tbl = Table(data, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            flow.append(tbl)
        else:
            flow.append(Paragraph(_pdf_inline(b["text"]), styles["BodyText"]))
        flow.append(Spacer(1, 6))
    SimpleDocTemplate(str(out), pagesize=A4).build(flow)


# --- output path + CLI -------------------------------------------------------

def _default_dir(src: Path) -> Path:
    for anc in [src.parent, *src.parents]:
        if (anc / ".cdt").is_dir():
            return anc / ".cdt" / "deliverables"
    return src.parent


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        prog="cdt doc",
        description="Render a Markdown file to a .docx and/or .pdf deliverable.",
    )
    parser.add_argument("source", help="path to the Markdown file")
    parser.add_argument("--out", metavar="FILE", help="output path (extension set "
                        "by --format; default: <project>/.cdt/deliverables/<name>.<ext>)")
    parser.add_argument("--format", choices=["docx", "pdf", "both"], default="docx",
                        help="output format (default: docx)")
    parser.add_argument("--title", help="document title")
    args = parser.parse_args(argv)

    src = Path(args.source).expanduser()
    if not src.is_file():
        print(f"cdt doc: file not found: {src}", file=sys.stderr)
        return 2

    fmts = ["docx", "pdf"] if args.format == "both" else [args.format]
    missing = []
    if "docx" in fmts:
        try:
            import docx  # noqa: F401
        except ImportError:
            missing.append("python-docx (.docx)")
    if "pdf" in fmts:
        try:
            import reportlab  # noqa: F401
        except ImportError:
            missing.append("reportlab (.pdf)")
    if missing:
        print("cdt doc: missing renderer(s): " + ", ".join(missing)
              + ". Run: pip install conductor[docs]", file=sys.stderr)
        return 2

    blocks = parse_blocks(src.read_text(encoding="utf-8"))
    out_base = Path(args.out).expanduser() if args.out else (_default_dir(src) / src.name)
    out_base.parent.mkdir(parents=True, exist_ok=True)
    for fmt in fmts:
        out = out_base.with_suffix("." + fmt)
        (render_docx if fmt == "docx" else render_pdf)(blocks, out, args.title)
        print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
